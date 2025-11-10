from fastapi import FastAPI, Request, Form, File, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pathlib import Path
from typing import Optional
from io import BytesIO
from docx import Document
from helpers.gemini_helper import rewrite_cover_letter

app = FastAPI(title="Cover Letter Tweaker")

# Create directories if they don't exist
Path("static").mkdir(exist_ok=True)
Path("templates").mkdir(exist_ok=True)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# Setup templates
templates = Jinja2Templates(directory="templates")


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """Render the main page"""
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/process")
async def process_cover_letter(
    companyDescription: str = Form(...),
    roleDescription: str = Form(...),
    resumeText: Optional[str] = Form(None),
    coverLetterText: Optional[str] = Form(None),
    resumeFile: Optional[UploadFile] = File(None),
    coverLetterFile: Optional[UploadFile] = File(None),
):
    """Process the cover letter using Gemini AI

    Combines company and role descriptions, then uses the rewrite_cover_letter
    function to generate a customized cover letter based on the resume.
    Accepts either text or file upload (TXT, PDF, DOCX) for both resume and cover letter.
    """
    try:
        # Validate that we have either resume text or file
        if not resumeText and not resumeFile:
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": "Please provide either resume text or upload a resume file.",
                },
            )

        # Validate that we have either cover letter text or file
        if not coverLetterText and not coverLetterFile:
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": "Please provide either cover letter text or upload a file.",
                },
            )

        # Combine company and role descriptions into job_details
        job_details = f"Company Description:\n{companyDescription}\n\nRole Description:\n{roleDescription}"

        # Handle resume file upload if provided
        resume_file_data = None
        resume_file_mime_type = None

        if resumeFile:
            # Read file bytes
            file_bytes = await resumeFile.read()

            # Validate file size (10MB limit for reasonable processing)
            max_size = 10 * 1024 * 1024  # 10MB
            if len(file_bytes) > max_size:
                return JSONResponse(
                    status_code=400,
                    content={
                        "success": False,
                        "error": "Resume file size exceeds 10MB limit. Please use a smaller file.",
                    },
                )

            # Determine file type and process accordingly
            filename = resumeFile.filename.lower()

            if filename.endswith(".txt"):
                # TXT: pass as inline data
                resume_file_mime_type = "text/plain"
                resume_file_data = file_bytes

            elif filename.endswith(".pdf"):
                # PDF: pass as inline data (Gemini supports this)
                resume_file_mime_type = "application/pdf"
                resume_file_data = file_bytes

            elif filename.endswith(".docx"):
                # DOCX: extract text on backend (Gemini doesn't support DOCX inline data)
                try:
                    doc = Document(BytesIO(file_bytes))
                    resumeText = "\n".join(
                        [paragraph.text for paragraph in doc.paragraphs]
                    )
                    # Clear file data so we use text path instead
                    resume_file_data = None
                    resume_file_mime_type = None
                except Exception as e:
                    return JSONResponse(
                        status_code=400,
                        content={
                            "success": False,
                            "error": f"Failed to extract text from resume DOCX file: {e}",
                        },
                    )
            else:
                return JSONResponse(
                    status_code=400,
                    content={
                        "success": False,
                        "error": "Unsupported resume file type. Please upload a .txt, .pdf, or .docx file.",
                    },
                )

        # Handle cover letter file upload if provided
        cover_letter_file_data = None
        cover_letter_file_mime_type = None

        if coverLetterFile:
            # Read file bytes
            file_bytes = await coverLetterFile.read()

            # Validate file size (10MB limit for reasonable processing)
            max_size = 10 * 1024 * 1024  # 10MB
            if len(file_bytes) > max_size:
                return JSONResponse(
                    status_code=400,
                    content={
                        "success": False,
                        "error": "Cover letter file size exceeds 10MB limit. Please use a smaller file.",
                    },
                )

            # Determine file type and process accordingly
            filename = coverLetterFile.filename.lower()

            if filename.endswith(".txt"):
                # TXT: pass as inline data
                cover_letter_file_mime_type = "text/plain"
                cover_letter_file_data = file_bytes

            elif filename.endswith(".pdf"):
                # PDF: pass as inline data (Gemini supports this)
                cover_letter_file_mime_type = "application/pdf"
                cover_letter_file_data = file_bytes

            elif filename.endswith(".docx"):
                # DOCX: extract text on backend (Gemini doesn't support DOCX inline data)
                try:
                    doc = Document(BytesIO(file_bytes))
                    coverLetterText = "\n".join(
                        [paragraph.text for paragraph in doc.paragraphs]
                    )
                    # Clear file data so we use text path instead
                    cover_letter_file_data = None
                    cover_letter_file_mime_type = None
                except Exception as e:
                    return JSONResponse(
                        status_code=400,
                        content={
                            "success": False,
                            "error": f"Failed to extract text from cover letter DOCX file: {e}",
                        },
                    )
            else:
                return JSONResponse(
                    status_code=400,
                    content={
                        "success": False,
                        "error": "Unsupported cover letter file type. Please upload a .txt, .pdf, or .docx file.",
                    },
                )

        # Call the rewrite_cover_letter function
        revised_letter = rewrite_cover_letter(
            job_details=job_details,
            resume_text=resumeText,
            existing_letter=coverLetterText,
            coverLetter_file_data=cover_letter_file_data,
            coverLetter_file_mime_type=cover_letter_file_mime_type,
            resume_file_data=resume_file_data,
            resume_file_mime_type=resume_file_mime_type,
        )

        if revised_letter:
            return JSONResponse(
                content={"success": True, "revised_letter": revised_letter}
            )
        else:
            return JSONResponse(
                status_code=500,
                content={
                    "success": False,
                    "error": "Failed to generate revised cover letter. Please check your API key and try again.",
                },
            )

    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": f"An error occurred: {str(e)}"},
        )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8000)
